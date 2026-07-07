from __future__ import annotations

import sys
from pathlib import Path
from typing import List, Tuple

try:
    import ollama
except ImportError:
    sys.exit("Missing dependency 'ollama'. Run: pip install ollama")

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx as docx_lib  # python-docx
except ImportError:
    docx_lib = None

from . import config

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp"}
TEXT_EXTENSIONS = {".txt", ".md"}
CODE_EXTENSIONS = {
    ".py", ".js", ".jsx", ".ts", ".tsx", ".java", ".c", ".cpp", ".h", ".hpp",
    ".cs", ".go", ".rs", ".rb", ".php", ".swift", ".kt", ".sql", ".sh", ".ps1",
    ".html", ".css", ".json", ".yaml", ".yml",
}


def load_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def load_code(path: Path) -> str:
    """Read source code as-is — no whitespace normalisation, since indentation
    and exact formatting matter for understanding/fixing code."""
    return path.read_text(encoding="utf-8", errors="ignore")


def load_pdf(path: Path) -> str:
    if PdfReader is None:
        raise RuntimeError("pypdf not installed. Run: pip install pypdf")
    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            parts.append("")
    return "\n".join(parts)


def load_docx(path: Path) -> str:
    if docx_lib is None:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")
    document = docx_lib.Document(str(path))
    paras = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paras.append(cell.text)
    return "\n".join(paras)


def describe_image(path: Path) -> str:
    """Use a vision model (llava) to turn an image into a text description
    so it can be embedded and retrieved like any other chunk."""
    try:
        resp = ollama.chat(
            model=config.VISION_MODEL,
            messages=[{
                "role": "user",
                "content": "Describe this image in detail, including any visible text, "
                            "charts, diagrams, or data. Be factual and thorough.",
                "images": [str(path)],
            }],
        )
        return resp["message"]["content"].strip()
    except Exception as e:  # noqa: BLE001
        raise RuntimeError(
            f"Vision model call failed (model={config.VISION_MODEL}). "
            f"Run 'ollama pull {config.VISION_MODEL}' for multimodal mode. Error: {e}"
        )


def load_document(path: Path) -> Tuple[str, str]:
    """Returns (raw_text, doc_type) for a given file path."""
    ext = path.suffix.lower()
    if ext in CODE_EXTENSIONS:
        return load_code(path), "code"
    if ext in TEXT_EXTENSIONS:
        return load_txt(path), "text"
    if ext == ".pdf":
        return load_pdf(path), "pdf"
    if ext == ".docx":
        return load_docx(path), "docx"
    if ext in IMAGE_EXTENSIONS:
        return describe_image(path), "image"
    raise ValueError(f"Unsupported file type: {ext} ({path.name})")


def discover_files(root: Path) -> List[Path]:
    """Recursively find ingestible files under a path (or return the single file)."""
    supported = TEXT_EXTENSIONS | {".pdf", ".docx"} | IMAGE_EXTENSIONS | CODE_EXTENSIONS
    if root.is_file():
        return [root] if root.suffix.lower() in supported else []
    files = []
    skip_dirs = {"node_modules", ".git", "__pycache__", "venv", ".venv", "dist", "build"}
    for p in root.rglob("*"):
        if any(part in skip_dirs for part in p.parts):
            continue
        if p.is_file() and p.suffix.lower() in supported:
            files.append(p)
    return sorted(files)
